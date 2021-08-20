# Import libraries
from xml.dom import minidom
import uuid
import docx
import re
import base64
import platform
import os
import tkinter as tk
from tkinter import filedialog
from PIL import ImageFont

# Global variables
gWidth = 1024                        # Screen width for main screen
gHeight = 768                        # Screen height for main screen
gMargin = 50                         # Margin for text
gMaxWidth = gWidth - (gMargin * 2)   # Max width = screen width - margin widths
gMaxHeight = gHeight - (gMargin * 2) # Max height = screen height - margin heights
gSansSize = gMargin * 1.5            # Sans size = 1.5 * margin (in this case, 75)
gGeorgiaSize = gMargin * 1.1         # Georgia size = 1.1 * margin (in this case, 55)
gIndent = "•  "                      # Bullet point
if platform.system() == "Windows":   # Specify font locations in Windows
    gCMGFontPath = os.environ["LOCALAPPDATA"] + "\Microsoft\Windows\Fonts\CMGSans-Thin.ttf"
    gGeorgiaFontPath = "C:\Windows\Fonts\georgia.ttf"
    gInitialDir = os.path.expanduser("~") + "\\Dropbox\\Media\\sermon notes\\"
    gSaveDir = os.path.expanduser("~") + "\\Documents\\ProPresenter6\\"
elif platform.system() == "Darwin":  # Specify font locations in Mac OS
    gCMGFontPath = os.path.expanduser("~") + "/Library/Fonts/CMGSans-Thin.ttf"
    gGeorgiaFontPath = "/System/Library/Fonts/Supplemental/Georgia.ttf"
    gInitialDir = os.path.expanduser("~") + "/Dropbox/Media/sermon notes/"
    gSaveDir = os.path.expanduser("~") + "/Documents/ProPresenter6/"
else:                                # Not intended for use on any other OS, as ProPresenter only runs on Windows and Mac
    raise Exception("This script is not intended for use outside of Windows or Mac operating systems.")

# Functions
def format_scripture(string):
    book = re.sub(r"(?s)\A((?:[1-3] *)?(?:Song of )?\w+\.?).+", r"\1", string)
    chapter = re.sub(r"(?:[1-3] *)?(?:Song of )?\w+\.? (\d+).*", r"\1", string)
    vnums = re.findall(r"[: ](\d+) ", string)
    if not len(vnums):
        raise Exception("No verse numbers were found.")
        return string
    if len(vnums) > 1:
        verses = ""
        i = 0
        while i <= len(vnums):
            if i == 0:
                verses += str(len(vnums))
            else:
                if vnums[1] - vnums[i - 1] == 1:
                    while vnums[i] - vnums[i - 1] == 1:
                        i += 1
                    verses += "-" + str(vnums[i])
                else:
                    verses += ", " + str(vnums[i])
            i += 1
    else:
        verses = vnums[0]
    value = (re.sub(r"((?:[1-3] *)?(?:Song of )?\w+\.? \d+:\d+|\d+)", "", string) + "\r\n" + "%s %s:%s" % (book, chapter, verses)).strip()
    return value


def generate_bible_slide(string, font_size=int(gGeorgiaSize)):
    font_path = gCMGFontPath
    imgfont = ImageFont.truetype(font=font_path, size=int(font_size - 5))
    refheight = imgfont.getsize("Hy")[1] * 1.25
    if not re.search(r"\r\n", string):
        raise Exception("String does not contain carriage return.")
        return

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide node
    RVDisplaySlide = root.createElement("RVDisplaySlide")
    RVDisplaySlide.setAttribute("UUID", str(uuid.uuid4()).upper())
    RVDisplaySlide.setAttribute("backgroundColor", "0 0 0 1")
    RVDisplaySlide.setAttribute("chordChartPath", "")
    RVDisplaySlide.setAttribute("drawingBackgroundColor", "false")
    RVDisplaySlide.setAttribute("enabled", "true")
    RVDisplaySlide.setAttribute("highlightColor", "0 0 0 0")
    RVDisplaySlide.setAttribute("hotKey", "")
    RVDisplaySlide.setAttribute("label", re.sub(r".*\r\n(.*?)", r"\1", string))
    RVDisplaySlide.setAttribute("notes", "")
    RVDisplaySlide.setAttribute("socialItemCount", "1")
    slides.appendChild(RVDisplaySlide)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array node (cues)
    empty = root.createTextNode("")
    cues = root.createElement("array")
    cues.setAttribute("rvXMLIvarName", "cues")
    RVDisplaySlide.appendChild(cues)
    cues.appendChild(empty)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array node (displayElements)
    displayElements = root.createElement("array")
    displayElements.setAttribute("rvXMLIvarName", "displayElements")
    RVDisplaySlide.appendChild(displayElements)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement node
    RVTextElement = root.createElement("RVTextElement")
    RVTextElement.setAttribute("UUID", str(uuid.uuid4()).upper())
    RVTextElement.setAttribute("additionalLineFillHeight", "0.000000")
    RVTextElement.setAttribute("adjustsHeightToFit", "false")
    RVTextElement.setAttribute("bezelRadius", "0.000000")
    RVTextElement.setAttribute("displayDelay", "0.000000")
    RVTextElement.setAttribute("displayName", "Bible Text")
    RVTextElement.setAttribute("drawLineBackground", "false")
    RVTextElement.setAttribute("drawingFill", "false")
    RVTextElement.setAttribute("drawingShadow", "true")
    RVTextElement.setAttribute("drawingStroke", "false")
    RVTextElement.setAttribute("fillColor", "0 0 0 0")
    RVTextElement.setAttribute("fromTemplate", "false")
    RVTextElement.setAttribute("lineBackgroundType", "0")
    RVTextElement.setAttribute("lineFillVerticalOffset", "0.000000")
    RVTextElement.setAttribute("locked", "false")
    RVTextElement.setAttribute("opacity", "1.000000")
    RVTextElement.setAttribute("persistent", "false")
    RVTextElement.setAttribute("revealType", "0")
    RVTextElement.setAttribute("rotation", "0.000000")
    RVTextElement.setAttribute("source", "")
    RVTextElement.setAttribute("textSourceRemoveLineReturnsOption", "false")
    RVTextElement.setAttribute("typeID", "0")
    RVTextElement.setAttribute("useAllCaps", "false")
    RVTextElement.setAttribute("verticalAlignment", "1")
    displayElements.appendChild(RVTextElement)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/RVRect3D node
    RVRect3D = root.createElement("RVRect3D")
    rectstring = "{%s %s 0 %s %s}" % (str(gMargin), str(gMargin), str(gMaxWidth), str(gMaxHeight - refheight))
    # print("Bible Text: " + rectstring)
    RVRect3Dstring = root.createTextNode(rectstring)
    RVRect3D.setAttribute("rvXMLIvarName", "position")
    RVTextElement.appendChild(RVRect3D)
    RVRect3D.appendChild(RVRect3Dstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/shadow node
    shadow = root.createElement("shadow")
    shadowstring = root.createTextNode(r"3.000000|0 0 0 1|{4, -4}")
    shadow.setAttribute("rvXMLIvarName", "shadow")
    RVTextElement.appendChild(shadow)
    shadow.appendChild(shadowstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary node (stroke)
    stroke = root.createElement("dictionary")
    stroke.setAttribute("rvXMLIvarName", "stroke")
    RVTextElement.appendChild(stroke)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary/NSColor node
    NSColor = root.createElement("NSColor")
    NSColorstring = root.createTextNode("0 0 0 1")
    NSColor.setAttribute("rvXMLDictionaryKey", "RVShapeElementStrokeColorKey")
    stroke.appendChild(NSColor)
    NSColor.appendChild(NSColorstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary/NSNumber node
    NSNumber = root.createElement("NSNumber")
    NSNumberstring = root.createTextNode("1.000000")
    NSNumber.setAttribute("hint", "float")
    NSNumber.setAttribute("rvXMLDictionaryKey", "RVShapeElementStrokeWidthKey")
    stroke.appendChild(NSNumber)
    NSNumber.appendChild(NSNumberstring)

    # Prepare RTF data for Bible Text
    rtfAppend = re.sub(r"(.*?)\r\n.*", r"\1", string.strip())
    replacements = [["•", r"\'95"],["“", r"\ldblquote "],["”", r"\rdblquote "],["‘", r"\lquote "],["’", r"\rquote "],["–", r"\'96"],["{", r"\{"],["}", r"\}"]]
    for replacement in replacements:
        if replacement[0] in rtfAppend:
            rtfAppend = rtfAppend.replace(replacement[0], replacement[1])
    rtf = r"{\rtf1\ansi\ansicpg1252\cocoartf2513" + "\r\n"
    rtf += r"\cocoatextscaling0\cocoaplatform0{\fonttbl\f0\fnil\fcharset0 Georgia;}" + "\r\n"
    rtf += r"{\colortbl;\red255\green255\blue255;\red255\green255\blue255;}" + "\r\n"
    rtf += r"{\*\expandedcolortbl;;\csgray\c100000;}" + "\r\n"
    rtf += r"\deftab720" + "\r\n"
    rtf += r"\pard\pardeftab720\qj\partightenfactor0" + "\r\n" + "\r\n"
    rtf += r"\f0\fs" + str(int(font_size * 2)) + r" \cf2 \outl0\strokewidth-40 \strokec0 " + rtfAppend.strip() + "}"
    # print(rtf)
    # print("=" * 80)
    rtf_enc = str(base64.b64encode(rtf.encode("utf8")), "utf-8")

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/NSString node
    NSString = root.createElement("NSString")
    NSStringstring = root.createTextNode(rtf_enc)
    NSString.setAttribute("rvXMLIvarName", "RTFData")
    RVTextElement.appendChild(NSString)
    NSString.appendChild(NSStringstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement node
    RVTextElement = root.createElement("RVTextElement")
    RVTextElement.setAttribute("UUID", str(uuid.uuid4()).upper())
    RVTextElement.setAttribute("additionalLineFillHeight", "0.000000")
    RVTextElement.setAttribute("adjustsHeightToFit", "false")
    RVTextElement.setAttribute("bezelRadius", "0.000000")
    RVTextElement.setAttribute("displayDelay", "0.000000")
    RVTextElement.setAttribute("displayName", "Bible Reference")
    RVTextElement.setAttribute("drawLineBackground", "false")
    RVTextElement.setAttribute("drawingFill", "false")
    RVTextElement.setAttribute("drawingShadow", "true")
    RVTextElement.setAttribute("drawingStroke", "false")
    RVTextElement.setAttribute("fillColor", "0 0 0 0")
    RVTextElement.setAttribute("fromTemplate", "true")
    RVTextElement.setAttribute("lineBackgroundType", "0")
    RVTextElement.setAttribute("lineFillVerticalOffset", "0.000000")
    RVTextElement.setAttribute("locked", "false")
    RVTextElement.setAttribute("opacity", "1.000000")
    RVTextElement.setAttribute("persistent", "false")
    RVTextElement.setAttribute("revealType", "0")
    RVTextElement.setAttribute("rotation", "0.000000")
    RVTextElement.setAttribute("source", "")
    RVTextElement.setAttribute("textSourceRemoveLineReturnsOption", "false")
    RVTextElement.setAttribute("typeID", "0")
    RVTextElement.setAttribute("useAllCaps", "false")
    RVTextElement.setAttribute("verticalAlignment", "1")
    displayElements.appendChild(RVTextElement)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/RVRect3D node
    RVRect3D = root.createElement("RVRect3D")
    rectstring = "{%s %s 0 %s %s}" % (str(gMargin), str(gHeight - gMargin - refheight), str(gMaxWidth), refheight)
    # print("Bible Reference: " + rectstring)
    RVRect3Dstring = root.createTextNode(rectstring)
    RVRect3D.setAttribute("rvXMLIvarName", "position")
    RVTextElement.appendChild(RVRect3D)
    RVRect3D.appendChild(RVRect3Dstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/shadow node
    shadow = root.createElement("shadow")
    shadowstring = root.createTextNode(r"3.000000|0 0 0 1|{4, -4}")
    shadow.setAttribute("rvXMLIvarName", "shadow")
    RVTextElement.appendChild(shadow)
    shadow.appendChild(shadowstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary node (stroke)
    stroke = root.createElement("dictionary")
    stroke.setAttribute("rvXMLIvarName", "stroke")
    RVTextElement.appendChild(stroke)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary/NSColor node
    NSColor = root.createElement("NSColor")
    NSColorstring = root.createTextNode("0 0 0 1")
    NSColor.setAttribute("rvXMLDictionaryKey", "RVShapeElementStrokeColorKey")
    stroke.appendChild(NSColor)
    NSColor.appendChild(NSColorstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary/NSNumber node
    NSNumber = root.createElement("NSNumber")
    NSNumberstring = root.createTextNode("1.000000")
    NSNumber.setAttribute("hint", "float")
    NSNumber.setAttribute("rvXMLDictionaryKey", "RVShapeElementStrokeWidthKey")
    stroke.appendChild(NSNumber)
    NSNumber.appendChild(NSNumberstring)

    # Prepare RTF data for Bible Reference
    rtfAppend = re.sub(r".*?\r\n(.*?)", r"\1", string).strip()
    for replacement in replacements:
        if replacement[0] in rtfAppend:
            rtfAppend = rtfAppend.replace(replacement[0], replacement[1])
    rtf = r"{\rtf1\ansi\ansicpg1252\cocoartf2513" + "\r\n"
    rtf += r"\cocoatextscaling0\cocoaplatform0{\fonttbl\f0\fnil\fcharset0 CMGSans-Regular;}" + "\r\n"
    rtf += r"{\colortbl;\red255\green255\blue255;\red255\green255\blue0;}" + "\r\n"
    rtf += r"{\*\expandedcolortbl;;\csgenericrgb\c100000\c100000\c0;}" + "\r\n"
    rtf += r"\pard\pardirnatural\qr\partightenfactor0" + "\r\n" + "\r\n"
    rtf += r"\f0\b\fs" + str(int((font_size - 5) * 2)) + r" \cf2 \out10\strokewidth-60 \strokec0 " + rtfAppend.strip() + "}"
    # print(rtf)
    # print("=" * 80)
    rtf_enc = str(base64.b64encode(rtf.encode("utf8")), "utf-8")

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/NSString node
    NSString = root.createElement("NSString")
    NSStringstring = root.createTextNode(rtf_enc)
    NSString.setAttribute("rvXMLIvarName", "RTFData")
    RVTextElement.appendChild(NSString)
    NSString.appendChild(NSStringstring)
    return


def generate_slide(string, font_size=int(gSansSize)):
    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide node
    RVDisplaySlide = root.createElement("RVDisplaySlide")
    RVDisplaySlide.setAttribute("UUID", str(uuid.uuid4()).upper())
    RVDisplaySlide.setAttribute("backgroundColor", "0 0 0 1")
    RVDisplaySlide.setAttribute("chordChartPath", "")
    RVDisplaySlide.setAttribute("drawingBackgroundColor", "false")
    RVDisplaySlide.setAttribute("enabled", "true")
    RVDisplaySlide.setAttribute("highlightColor", "0 0 0 0")
    RVDisplaySlide.setAttribute("hotKey", "")
    RVDisplaySlide.setAttribute("label", "")
    RVDisplaySlide.setAttribute("notes", "")
    RVDisplaySlide.setAttribute("socialItemCount", "1")
    slides.appendChild(RVDisplaySlide)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array node (cues)
    empty = root.createTextNode("")
    cues = root.createElement("array")
    cues.setAttribute("rvXMLIvarName", "cues")
    RVDisplaySlide.appendChild(cues)
    cues.appendChild(empty)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array node (displayElements)
    displayElements = root.createElement("array")
    displayElements.setAttribute("rvXMLIvarName", "displayElements")
    RVDisplaySlide.appendChild(displayElements)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement node
    RVTextElement = root.createElement("RVTextElement")
    RVTextElement.setAttribute("UUID", str(uuid.uuid4()).upper())
    RVTextElement.setAttribute("additionalLineFillHeight", "0.000000")
    RVTextElement.setAttribute("adjustsHeightToFit", "false")
    RVTextElement.setAttribute("bezelRadius", "0.000000")
    RVTextElement.setAttribute("displayDelay", "0.000000")
    RVTextElement.setAttribute("displayName", "TextElement")
    RVTextElement.setAttribute("drawLineBackground", "false")
    RVTextElement.setAttribute("drawingFill", "false")
    RVTextElement.setAttribute("drawingShadow", "true")
    RVTextElement.setAttribute("drawingStroke", "false")
    RVTextElement.setAttribute("fillColor", "0 0 0 0")
    RVTextElement.setAttribute("fromTemplate", "true")
    RVTextElement.setAttribute("lineBackgroundType", "0")
    RVTextElement.setAttribute("lineFillVerticalOffset", "0.000000")
    RVTextElement.setAttribute("locked", "false")
    RVTextElement.setAttribute("opacity", "1.000000")
    RVTextElement.setAttribute("persistent", "false")
    RVTextElement.setAttribute("revealType", "0")
    RVTextElement.setAttribute("rotation", "0.000000")
    RVTextElement.setAttribute("source", "")
    RVTextElement.setAttribute("textSourceRemoveLineReturnsOption", "false")
    RVTextElement.setAttribute("typeID", "0")
    RVTextElement.setAttribute("useAllCaps", "false")
    RVTextElement.setAttribute("verticalAlignment", "1")
    displayElements.appendChild(RVTextElement)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/RVRect3D node
    RVRect3D = root.createElement("RVRect3D")
    rectstring = root.createTextNode("{50 50 0 924 668}")
    RVRect3D.setAttribute("rvXMLIvarName", "position")
    RVTextElement.appendChild(RVRect3D)
    RVRect3D.appendChild(rectstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/shadow node
    shadow = root.createElement("shadow")
    shadowstring = root.createTextNode("0.000000|0 0 0 1|{4, -4}")
    shadow.setAttribute("rvXMLIvarName", "shadow")
    RVTextElement.appendChild(shadow)
    shadow.appendChild(shadowstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary node (stroke)
    stroke = root.createElement("dictionary")
    stroke.setAttribute("rvXMLIvarName", "stroke")
    RVTextElement.appendChild(stroke)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary/NSColor node
    NSColor = root.createElement("NSColor")
    NSColorstring = root.createTextNode("0 0 0 1")
    NSColor.setAttribute("rvXMLDictionaryKey", "RVShapeElementStrokeColorKey")
    stroke.appendChild(NSColor)
    NSColor.appendChild(NSColorstring)

    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/dictionary/NSNumber node
    NSNumber = root.createElement("NSNumber")
    NSNumberstring = root.createTextNode("0.000000")
    NSNumber.setAttribute("hint", "double")
    NSNumber.setAttribute("rvXMLDictionaryKey", "RVShapeElementStrokeWidthKey")
    stroke.appendChild(NSNumber)
    NSNumber.appendChild(NSNumberstring)

    # Prepare RTF data
    pattern = r"((?:[1-3] *)?(?:Song of )?\w+ \d+\:[0-9;,-:]+|(?:\b[Vv]\w*?\.? )?[0-9-,]+\Z)"
    rtfAppend = string.strip()
    if re.search(r"\r\n", rtfAppend):
        rtfAppend = re.sub(r"(\r\n+)", r"\\\1", rtfAppend)
    if re.search(pattern, string):
        rtfAppend = re.sub(pattern, r"{\\cf2 \1}", rtfAppend)
    replacements = [["•", r"\'95"],["“", r"\ldblquote "],["”", r"\rdblquote "],["‘", r"\lquote "],["’", r"\rquote "],["–", r"\'96"],["\t", r"\tab "]]
    for replacement in replacements:
        if replacement[0] in rtfAppend:
            rtfAppend = rtfAppend.replace(replacement[0], replacement[1])
    rtf = r"{\rtf1\ansi\ansicpg1252\cocoartf1671\cocoasubrtf600" + "\r\n"
    rtf += r"{\fonttbl\f0\fnil\fcharset0 CMGSans-Regular;}" + "\r\n"
    rtf += r"{\colortbl;\red255\green255\blue255;"
    if re.search(pattern, string):
        rtf += r"\red255\green255\blue0;"
    rtf += "}" + "\r\n"
    rtf += r"\pard\pardirnatural"
    if "•" in string:
        rtf += r"\ql"
    else:
        rtf += r"\qc"
    rtf += r"\partightenfactor0" + "\r\n" + "\r\n"
    rtf += r"\f0\b\fs" + str(int(font_size * 2)) + r" \cf1 \outl0\strokewidth-60 \strokec0 " + rtfAppend.strip() + "}"
    # print(rtf)
    # print("=" * 80)
    rtf_enc = str(base64.b64encode(rtf.encode("utf8")), "utf-8")
    
    # Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array/RVDisplaySlide/array/RVTextElement/NSString node
    NSString = root.createElement("NSString")
    NSStringstring = root.createTextNode(rtf_enc)
    NSString.setAttribute("rvXMLIvarName", "RTFData")
    RVTextElement.appendChild(NSString)
    NSString.appendChild(NSStringstring)
    return


def parseIntSet(nputstr):
    selection = set()
    invalid = set()
    # tokens are comma seperated values
    tokens = [x.strip() for x in nputstr.split(",")]
    for i in tokens:
        try:
            # typically tokens are plain old integers
            selection.add(int(i))
        except:
            # if not, then it might be a range
            try:
                token = [int(k.strip()) for k in i.split("-")]
                if len(token) > 1:
                    token.sort()
                    # we have items seperated by a dash
                    # try to build a valid range
                    first = token[0]
                    last = token[len(token)-1]
                    for x in range(first, last+1):
                        selection.add(x)
            except:
                # not an int and not a range...
                invalid.add(i)
    # Report invalid tokens before returning valid selection
    if len(invalid) > 0: print("Invalid set: " + str(invalid))
    return selection


def string_wrap(string, font="CMG", font_size=int(gSansSize), max_width=int(gMaxWidth), max_height=int(gMaxHeight)):
    if string.strip() == "":
        return (int(font_size), string.strip())
    if font == "CMG":
        font_path = gCMGFontPath
    elif font == "Georgia":
        font_path = gGeorgiaFontPath
    else:
        raise Exception("Font path has not been specified for " + font + ".")
        return (int(font_size), string.strip())
    imgfont = ImageFont.truetype(font=font_path, size=int(font_size))

    subindent = ""
    if gIndent in string:
        indentwidth = imgfont.getsize(gIndent)[0]
        while imgfont.getsize(subindent)[0] < indentwidth:
            subindent += " "
            if imgfont.getsize(subindent)[0] >= indentwidth: break

    words = string.split(" ")
    longest = ""
    for word in words:
        if len(word) > len(longest): longest = word
    if re.search(r"(?i)\A((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+))(?:[-–])?", string):
        scriptures = re.findall(r"(?i)((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+)(?:[-–])?)", string)
        for scripture in scriptures:
            if len(scripture) > len(longest): longest = scripture
    width = imgfont.getsize(subindent + longest)[0]
    while width > max_width:
        font_size -= 1
        # print("Decreased to " + str(font_size))
        imgfont = ImageFont.truetype(font=font_path, size=int(font_size))
        width = imgfont.getsize(subindent + longest)[0]

    lines = []

    wrapped = string
    width = imgfont.getsize(wrapped)[0]
    height = imgfont.getsize(wrapped)[1]
    if width <= max_width and height <= max_height:
        return (int(font_size), wrapped.strip())
    while width > max_width or height > max_height:
        words = wrapped.split(" ")
        i = 0
        while i < len(words):
            line = ""
            while i < len(words) and imgfont.getsize(subindent + line + words[i])[0] <= max_width:
                if len(lines) > 0 and line == "": line += subindent
                line += words[i]+ " "
                i += 1
            if not line:
                line = words[i]
                i += 1
            lines.append(line)
        width = 0
        height = 0
        for line in lines:
            sizes = imgfont.getsize(line)
            if sizes[0] > width: width = sizes[0]
            height += sizes[1]
        wrapped = "\r\n".join(lines)
        if width > max_width or height > max_height:
            font_size -= 1
            imgfont = ImageFont.truetype(font=font_path, size=int(font_size))
            wrapped = string
            lines = []
            if len(subindent):
                indentwidth = imgfont.getsize(indent)[0]
                while imgfont.getsize(subindent)[0] < indentwidth:
                    subindent += " "
                    if imgfont.getsize(subindent)[0] >= indentwidth: break
    if len(subindent) > 0:
        return (int(font_size), wrapped.strip())
    elif re.search(r"(?i)\A((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+) [-–])", string):
        string = re.sub(r"(?i)\A((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+) [-–]) (.*)", r"\1" + "\r\n" + r"\2", string)
        return (font_size, string.strip())
    elif re.search(r"(?i)((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+))\Z", string):
        string = re.sub(r"(?i)((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+))\Z", "\r\n" + r"\1", string)
        return (font_size, string.strip())
    else:
        return (int(font_size), string.strip())


# Create XML document
root = minidom.Document()

# Add RVPresentationDocument node
RVPresentationDocument = root.createElement("RVPresentationDocument")
RVPresentationDocument.setAttribute("CCLIArtistCredits", "")
RVPresentationDocument.setAttribute("CCLIAuthor", "")
RVPresentationDocument.setAttribute("CCLICopyrightYear", "")
RVPresentationDocument.setAttribute("CCLIDisplay", "false")
RVPresentationDocument.setAttribute("CCLIPublisher", "")
RVPresentationDocument.setAttribute("CCLISongNumber", "")
RVPresentationDocument.setAttribute("CCLISongTitle", "")
RVPresentationDocument.setAttribute("backgroundColor", "0 0 0 0")
RVPresentationDocument.setAttribute("buildNumber", "100991490")
RVPresentationDocument.setAttribute("category", "Sermon Notes")
RVPresentationDocument.setAttribute("chordChartPath", "")
RVPresentationDocument.setAttribute("docType", "0")
RVPresentationDocument.setAttribute("drawingBackgroundColor", "false")
RVPresentationDocument.setAttribute("height", "768")
RVPresentationDocument.setAttribute("lastDateUsed", "")
RVPresentationDocument.setAttribute("notes", "")
RVPresentationDocument.setAttribute("os", "2")
RVPresentationDocument.setAttribute("resourcesDirectory", "")
RVPresentationDocument.setAttribute("selectedArrangementID", "")
RVPresentationDocument.setAttribute("usedCount", "0")
RVPresentationDocument.setAttribute("uuid", str(uuid.uuid4()).upper())
RVPresentationDocument.setAttribute("versionNumber", "600")
RVPresentationDocument.setAttribute("width", "1024")
root.appendChild(RVPresentationDocument)

# Add RVPresentationDocument/RVTimeline node
RVTimeline = root.createElement("RVTimeline")
RVTimeline.setAttribute("duration", "0.000000")
RVTimeline.setAttribute("loop", "false")
RVTimeline.setAttribute("playBackRate", "1.000000")
RVTimeline.setAttribute("rvXMLIVarName", "timeline")
RVTimeline.setAttribute("selectedMediaTrackIndex", "-1")
RVTimeline.setAttribute("timeOffset", "0.000000")
RVPresentationDocument.appendChild(RVTimeline)

# Add RVPresentationDocument/RVTimeline/array node (timeCues)
empty = root.createTextNode("")
timeCues = root.createElement("array")
timeCues.setAttribute("rvXMLIvarName", "timeCues")
RVTimeline.appendChild(timeCues)
timeCues.appendChild(empty)

# Add RVPresentationDocument/RVTimeline/array node (mediaTracks)
empty = root.createTextNode("")
mediaTracks = root.createElement("array")
mediaTracks.setAttribute("rvXMLIvarName", "mediaTracks")
RVTimeline.appendChild(mediaTracks)
mediaTracks.appendChild(empty)

# Add RVPresentationDocument/array node (groups)
groups = root.createElement("array")
groups.setAttribute("rvXMLIvarName", "groups")
RVPresentationDocument.appendChild(groups)

# Add RVPresentationDocument/RVtimeline/array/RVSlideGrouping node
RVSlideGrouping = root.createElement("RVSlideGrouping")
RVSlideGrouping.setAttribute("color", "0 0 0 0")
RVSlideGrouping.setAttribute("name", "")
RVSlideGrouping.setAttribute("uuid", str(uuid.uuid4()).upper())
groups.appendChild(RVSlideGrouping)

# Add RVPresentationDocument/RVTimelne/array/RVSlideGrouping/array node (slides)
slides = root.createElement("array")
slides.setAttribute("rvXMLIvarName", "slides")
RVSlideGrouping.appendChild(slides)

# Get Word doc
tkroot = tk.Tk()
tkroot.withdraw()
docxfile = filedialog.askopenfile(title="Please select a Word document", initialdir=gInitialDir, filetypes=[("Word Documents", "*.docx")])
if docxfile == None: raise Exception("No file specified.")

print("Converting " + docxfile.name + "...")
document = docx.Document(docxfile.name)
lines = []
for paragraph in document.paragraphs: # Loop through paragraphs
    text = paragraph.text.strip()
    if paragraph.style.name == "List Paragraph": text = gIndent + text
    if text != "" and text not in lines: # Only do something if paragraph is not blank and isn't a duplicate
        lines.append(text)
lines[0] = ""
print(str(len(lines)) + " lines found")
linenbr = 1
for line in lines:
    if re.search(r"(?i)\A((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+)) (?!speaks|says|[-–])", line):
        line = format_scripture(line)
        versen = re.sub(r"(.*?)\r\n.*", r"\1", line)
        wrapped = string_wrap(versen, "Georgia", 55)
        # print("Size: " + str(wrapped[0]))
        generate_bible_slide(line, wrapped[0])
    elif re.search(r"(?i)\A((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+) [-–])", line):
        line = re.sub(r"(?i)\A((?:[1-3] *)?(?:Song of )?\w+\.? \d+\:(?:[0-9,:;-]+) [-–]) (.*)", r"\1" + "\r\n" + r"\2", line)
        wrapped = string_wrap(line)
        generate_slide(line, wrapped[0])
    else:
        wrapped = string_wrap(line)
        line = wrapped[1]
        generate_slide(line, wrapped[0])
    linenbr += 1

xml_str = root.childNodes[0].toxml()

save_file = os.path.basename(docxfile.name).replace(".docx", ".pro6")
save_path_file = filedialog.asksaveasfile(title="Save ProPresenter file", initialdir=gSaveDir, initialfile=save_file, filetype=[("ProPresenter Document", "*.pro6")])

with open(save_path_file.name, "w") as f:
    f.write(xml_str)

print("Finished!")
